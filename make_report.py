# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '주식투자방향_정리.docx')

doc = Document()

# ── 페이지 여백 ───────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 ──────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, font_name='맑은 고딕'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading(text, level=1, color=(0,70,127)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size = {1:16, 2:13, 3:11}[level]
    set_font(run, size=size, bold=True, color=color)
    return p

def add_body(text, indent=0, bullet=None, bold_parts=None, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    full = (bullet + ' ' + text) if bullet else text
    if bold_parts:
        remaining = full
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                set_font(r)
            r2 = p.add_run(bp)
            set_font(r2, bold=True)
            remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining)
            set_font(r)
    else:
        run = p.add_run(full)
        set_font(run)
    return p

def add_blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, header_color='1F4E79'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255,255,255))

    # 데이터
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        bg = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=9.5)

    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_box(lines, bg='F0F7FF', border_color='2E75B6'):
    """강조 박스"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), border_color.replace('#',''))
        tcBorders.append(border)
    tcPr.append(tcBorders)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(line)
        set_font(run, size=10)

# ═══════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(10)
run = p.add_run('주식투자 방향 정리')
set_font(run, size=26, bold=True, color=(31,78,121))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('월봉 10이평선 매매시스템 — 시뮬레이션 결과 기반')
set_font(run2, size=13, color=(89,89,89))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('2026년 5월')
set_font(run3, size=11, color=(128,128,128))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 목차
# ═══════════════════════════════════════════════════════════
add_heading('목  차', 1)
toc_items = [
    ('1', '시뮬레이션 결과 핵심 요약'),
    ('2', '계좌별 최적 투자 전략'),
    ('3', '현재 즉시 실행 사항 (2026년 5월 기준)'),
    ('4', '월봉 10이평 전략을 써야 하는 이유'),
    ('5', '장기 투자 원칙 5가지'),
    ('6', '자동화 알림 시스템 운용 방법'),
]
for num, title in toc_items:
    add_body(f'{num}.  {title}', indent=0.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. 시뮬레이션 결과 핵심 요약
# ═══════════════════════════════════════════════════════════
add_heading('1. 시뮬레이션 결과 핵심 요약', 1)
add_body('총 4가지 시뮬레이션을 통해 매매전략의 효과를 검증하였다.')
add_blank()

add_table(
    headers=['구분', '전략 수익률', 'B&H 수익률', '평균 MDD', '결론'],
    rows=[
        ['개별종목 24개\n(10년)', '+737%', '+2,225%', '-48%', 'B&H 우세\n전략은 리스크 관리'],
        ['미국 광의 ETF\n(QQQ·SPY 등, 20년)', '+340~759%', '+641~1,795%', '-19~40%', 'B&H 압도적 우세'],
        ['한국 테마 ETF\n(DC/IRP/연금)', '+57%', '+98%', '-15%', '2차전지만 전략 우위\n(데이터 짧음)'],
        ['복합 하이브리드\n(월봉+일봉, 10년)', '-1%', '—', '-12%', '강세주 진입 못함\nMDD만 낮음'],
    ],
    col_widths=[3.8, 2.8, 2.8, 2.5, 4.0],
)
add_blank()

add_heading('주요 개별종목 성과 (단순 월봉 10MA 전략)', 2)
add_table(
    headers=['종목', '이름', '전략 수익률', 'B&H 수익률', '비고'],
    rows=[
        ['NVDA', '엔비디아', '+3,232%', '+8,120%', '최대 성과 종목'],
        ['VRT',  '버티브',   '+5,227%', '+3,539%', '전략이 B&H 초과'],
        ['AMD',  'AMD',      '+1,890%', '+2,981%', '고수익 종목'],
        ['AVGO', '브로드컴', '+704%',   '+2,372%', '장기 우량주'],
        ['ASTS', 'AST스페이스','-15%',  '+622%',   '변동성 주의'],
        ['PLUG', '플러그파워','+643%',  '+158%',   '전략이 B&H 초과'],
    ],
    col_widths=[2.0, 3.0, 3.0, 3.0, 4.8],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. 계좌별 최적 투자 전략
# ═══════════════════════════════════════════════════════════
add_heading('2. 계좌별 최적 투자 전략', 1)

add_heading('2-1. DC / IRP / 연금저축 (한국 ETF)', 2)

add_heading('① 지수형 ETF — 적립식 B&H 유지', 3, color=(31,120,180))
add_body('KODEX 미국S&P500, TDF2045/2050, 채권혼합형 ETF', indent=0.5)
add_box([
    '원칙: 절대 매매하지 않는다.',
    '장기 복리가 목적인 계좌로, 매매할수록 손해다.',
    '매월 자동이체로 적립 후 방치한다.',
])
add_blank()

add_heading('② 테마형 ETF — 월봉 10이평 이탈 시 CD금리로 교체', 3, color=(31,120,180))
add_body('AI·반도체·빅테크·2차전지·원자력·테슬라 관련 ETF', indent=0.5)
add_blank()
add_table(
    headers=['신호', '행동', '이유'],
    rows=[
        ['10이평 위 → 홀딩', '현 ETF 그대로 유지', '추세 진행 중'],
        ['이번달 신규 돌파', 'CD금리 → 테마 ETF 교체', '재진입 신호'],
        ['이번달 이탈',     '테마 ETF → CD금리 교체', '손실 차단'],
        ['10이평 아래 유지', 'CD금리 계속 보유',       '추세 하락 중'],
    ],
    col_widths=[4.0, 4.5, 5.3],
    header_color='2E75B6',
)
add_blank()
add_body('★ 핵심: TIGER CD금리투자KIS(357870)가 피난처 역할. 연 3~4% 이자 수익.',
         bold_parts=['★ 핵심:', 'TIGER CD금리투자KIS(357870)'])

add_blank()
add_heading('2-2. 일반계좌 (미국 개별종목)', 2)
add_body('전략: 단순 월봉 10이평선 — 복합 하이브리드보다 효과적')
add_blank()

add_table(
    headers=['신호', '행동', '기준'],
    rows=[
        ['10이평 상향 돌파', '월말 종가 확정 후 매수 진입', '이격도 무관'],
        ['10이평 위 유지',   '보유 (홀딩)',                  '매매 없음'],
        ['10이평 하향 이탈', '즉시 전량 매도 (예외 없음)',   '이격도 무관'],
        ['10이평 아래 유지', '매수 절대 금지',               '반등해도 진입 안 함'],
    ],
    col_widths=[3.8, 5.5, 4.5],
    header_color='1F4E79',
)
add_blank()
add_body('복합 하이브리드 전략(월봉+일봉)은 진입 조건이 너무 엄격하여 NVDA·VRT 같은 '
         '강세주 진입 자체를 못한다. 단순 10MA 고수를 권장한다.',
         bold_parts=['단순 10MA 고수를 권장한다.'])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. 현재 즉시 실행 사항
# ═══════════════════════════════════════════════════════════
add_heading('3. 현재 즉시 실행 사항 (2026년 5월 기준)', 1)

add_heading('■ 즉시 매도 (10이평 이탈 확정)', 2, color=(192,0,0))
add_table(
    headers=['종목', '이름', '이격도', '조치'],
    rows=[
        ['LHX', 'L3해리스테크', '-0.5%', '즉시 전량 매도'],
        ['APH', '암페놀',       '-4.3%', '즉시 전량 매도'],
    ],
    col_widths=[2.0, 4.0, 3.0, 4.8],
    header_color='C00000',
)
add_blank()

add_heading('■ 매수 금지 (10이평 아래)', 2, color=(192,0,0))
add_table(
    headers=['종목', '이름', '이격도', '조치'],
    rows=[
        ['RTX',  'RTX',      '-1.0%',  '반등해도 진입 금지 — 돌파 확인 후 재진입'],
        ['LDOS', '레이도스', '-25.4%', '급락 경보 — 절대 매수 금지'],
    ],
    col_widths=[2.0, 3.0, 2.5, 6.3],
    header_color='C00000',
)
add_blank()

add_heading('■ 신규 진입 검토', 2, color=(31,120,50))
add_table(
    headers=['종목', '이름', '상태', '이격도', '조치'],
    rows=[
        ['IONQ', '아이온큐',    '★ 신규 돌파', '+19.1%', '월말 종가 확정 후 진입'],
        ['ASTS', 'AST스페이스', '▲ 눌림목',    '+0.3%',  '월말 확인 후 소량 진입 검토'],
    ],
    col_widths=[2.0, 3.0, 3.0, 2.5, 4.3],
    header_color='375623',
)
add_blank()

add_heading('■ DC/IRP/연금 테마 ETF 현황', 2, color=(31,78,121))
add_body('전체 10개 테마 ETF 모두 10이평 위 → 전량 홀딩 유지', indent=0.5)
add_body('SOL 미국원자력SMR: 이번 달 신규 돌파 → 이미 보유 중이면 유지', indent=0.5,
         bold_parts=['SOL 미국원자력SMR:'])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. 전략을 써야 하는 이유
# ═══════════════════════════════════════════════════════════
add_heading('4. 월봉 10이평 전략을 써야 하는 이유', 1)
add_body('시뮬레이션에서 B&H가 항상 높게 나온다. 그러나 B&H는 실제로 실행하기 어렵다.')
add_blank()

add_table(
    headers=['항목', 'B&H (단순보유)', '월봉 10MA 전략'],
    rows=[
        ['평균 수익률',   '+1,167%',      '+646%'],
        ['평균 MDD',     '-48% 이상',     '-20~30%'],
        ['심리적 고통',   '극심함',        '관리 가능'],
        ['실제 실행 가능','대부분 실패',   '규칙이 명확'],
        ['2022년 NVDA',  '-66% 견뎌야',   '이탈 시 매도'],
        ['2021~22 ARKK', '-75% 견뎌야',   '이탈 시 매도'],
    ],
    col_widths=[4.0, 5.0, 5.0],
)
add_blank()

add_box([
    '핵심 메시지:',
    '',
    '전략의 목적은 수익률 극대화가 아니다.',
    '시장에서 퇴출당하지 않고 꾸준히 복리를 쌓는 것이 목적이다.',
    '',
    '한 번의 -70% 폭락을 버티지 못하고 바닥에서 팔면',
    '이전 10년 수익이 모두 사라진다.',
    '전략은 그 최악의 순간을 막아준다.',
], bg='FFF2CC', border_color='FFB800')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. 장기 투자 원칙
# ═══════════════════════════════════════════════════════════
add_heading('5. 장기 투자 원칙 5가지', 1)

principles = [
    ('① 계좌 분리 원칙',
     ['DC/IRP/연금 → ETF 적립식 B&H + 테마 10이평 교체',
      '일반계좌 → 개별종목 단순 월봉 10MA',
      '두 계좌의 전략을 섞지 않는다.']),
    ('② 월 1회 확인 원칙',
     ['매달 말일 Slack 자동 알림 확인 → 즉시 실행',
      '중간에 차트를 자주 보면 감정적 판단이 생긴다.',
      '월봉이 기준이므로 월 1회 확인으로 충분하다.']),
    ('③ 감정적 판단 금지 원칙',
     ['10이평 이탈 → 뉴스·전망 무관하게 즉시 매도',
      '10이평 위 → 얼마나 올랐든 보유 유지',
      '"이번엔 다를 것 같다"는 생각이 가장 위험하다.']),
    ('④ 종목 수 관리 원칙',
     ['24개 종목 전부 보유할 필요 없다.',
      '이탈 시 매도 → 현금화 → 강한 종목만 자연스럽게 남는다.',
      '10개 이내 집중 보유가 관리하기 쉽고 수익도 높다.']),
    ('⑤ 적립은 ETF, 승부는 개별주 원칙',
     ['DC/IRP/연금 월 적립 → 복리 기반 장기 자산 형성',
      '일반계좌 개별주 → 알파 수익 추구',
      '두 역할을 명확히 구분한다.']),
]

for title, items in principles:
    add_heading(title, 2, color=(31,78,121))
    for item in items:
        add_body(item, indent=0.5, bullet='•')
    add_blank()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. 자동화 알림 시스템
# ═══════════════════════════════════════════════════════════
add_heading('6. 자동화 알림 시스템 운용 방법', 1)
add_body('모든 점검이 Slack으로 자동 발송된다. 별도 분석 없이 알림만 확인하면 된다.')
add_blank()

add_table(
    headers=['알림 종류', '발송 시점', '내용', '행동'],
    rows=[
        ['주간 모니터링', '매주 금요일\n오후 4시', '신규 돌파·이탈 경보\n눌림목 후보 목록', '관찰만 (매매 금지)'],
        ['월말 매매 결정', '매달 28~31일\n오후 4시 10분', '매수 진입·즉시 매도\n홀딩 목록 확정', '즉시 실행'],
        ['테마 ETF 점검', '위 알림에 포함', 'DC/IRP 테마 ETF\n10이평 위/아래 현황', 'CD금리 교체 여부 결정'],
    ],
    col_widths=[3.0, 3.0, 4.5, 3.3],
)
add_blank()

add_heading('알림 파일 경로', 2)
add_box([
    '주식정보 폴더 (c:\\Users\\user\\Desktop\\주식정보\\)',
    '',
    '  • slack_alert.py   — 알림 메인 스크립트',
    '  • config.json      — 종목 목록 및 설정',
    '  • scan_stocks.py   — 수동 스캔 (즉시 확인용)',
    '  • simulation.py    — 개별종목 백테스트',
    '  • etf_simulation.py        — 미국 ETF 백테스트',
    '  • korean_etf_simulation.py — 한국 ETF 백테스트',
    '  • hybrid_simulation.py     — 복합 하이브리드 백테스트',
    '',
    '  작업 스케줄러: Stock_Weekly_Friday / Stock_Monthly_28~31',
], bg='F5F5F5', border_color='AAAAAA')
add_blank()

add_heading('수동 알림 실행 방법', 2)
add_box([
    '즉시 테스트:  python slack_alert.py test',
    '주간 강제:    python slack_alert.py weekly',
    '월말 강제:    python slack_alert.py monthly',
], bg='EEF4FF', border_color='2E75B6')

add_blank()
add_blank()

# 마무리 문구
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('원칙대로 실행하면 된다.')
set_font(run, size=14, bold=True, color=(31,78,121))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('10이평 위 → 보유 / 10이평 이탈 → 즉시 매도 / 감정적 판단 금지')
set_font(run2, size=11, color=(89,89,89))

doc.save(OUT)
print(f'저장 완료: {OUT}')
